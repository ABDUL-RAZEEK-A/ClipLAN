import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Add a title
title = doc.add_heading('Google Play Store Publishing - Team Checklist', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add team details
doc.add_paragraph('Team: BATCH -1')
doc.add_paragraph('App name: ClipLAN: local share')
doc.add_paragraph('')

def add_table_section(heading, data):
    doc.add_heading(heading, level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Checklist Item'
    hdr_cells[1].text = 'Details / Answers'
    
    # Make headers bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                
    for key, value in data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value
    doc.add_paragraph('')

section1 = {
    "Team name / team number": "BATCH -1",
    "Team members' names": "Abdul Razeek A, Akash R, Aaron Michael Raj, Akash N",
    "App name (maximum 30 characters)": "ClipLAN: local share",
    "Default app language": "English (en-US)",
    "App or Game": "App",
    "Free or Paid": "Free",
    "App category": "Productivity / Tools",
    "Tags, if applicable": "File sharing, offline, peer-to-peer, transfer, productivity",
    "Package name / Application ID": "com.cliplan.cliplan",
    "Contact email": "abdulrazeek2006@gmail.com"
}
add_table_section('1. Basic App Details', section1)

section2 = {
    "Short description (maximum 80 characters)": "Blazing-fast, offline peer-to-peer file and clipboard sharing across devices.",
    "Full description (maximum 4,000 characters)": "ClipLAN is a blazing-fast, decentralized, offline, peer-to-peer (P2P) file and clipboard sharing application built to solve the bottleneck of transferring large files or text snippets between devices over local networks. The app is completely serverless, meaning it relies solely on the Local Area Network (LAN) or a mobile hotspot to communicate. By eliminating the cloud from the equation, ClipLAN achieves data transfer speeds limited only by the hardware limitations of the network interface cards and Wi-Fi routers involved.",
    "High-quality app icon": "Provided (app_logo.jpeg)",
    "Phone screenshots": "Provided (8 screenshots in launches/playstore/screenshots/final)",
    "Tablet screenshots, if applicable": "N/A",
    "Feature graphic, where required": "Provided",
    "Promotional video, if applicable": "N/A",
    "Developer/contact information": "Abdul Razeek A / abdulrazeek2006@gmail.com"
}
add_table_section('2. Store Listing Details', section2)

section3 = {
    "Final Android App Bundle (.aab)": "cliplan-release.aab (Provided)",
    "Application/package ID matches the submitted app": "Yes",
    "Final production build tested": "Yes",
    "App launches and core features work correctly": "Yes",
    "No unnecessary test/demo data in the production build": "Yes"
}
add_table_section('3. App Build & Technical Details', section3)

section4 = {
    "Privacy Policy URL": "Provided via in-app Modal & Website footer link",
    "Does the app collect user data? — Yes / No": "No",
    "Does the app share user data? — Yes / No": "No",
    "Types of data collected": "N/A",
    "Purpose of data collection": "N/A",
    "Third-party SDKs/services used": "None that collect or transmit user data. Local usage only.",
    "Is data encrypted in transit?": "Yes (Operates exclusively on local Wi-Fi / Hotspot networks, no cloud transit).",
    "Can users request deletion of their data?": "N/A (No data stored on external servers)."
}
add_table_section('4. Privacy & Data Safety', section4)

section5 = {
    "Contains ads? — Yes / No": "No",
    "Target audience / age group": "3+ / Everyone",
    "Content rating questionnaire information": "Everyone",
    "App access requirements": "Requires Local Network access (Wi-Fi) and Local Storage permissions.",
    "Sensitive permissions used": "MANAGE_EXTERNAL_STORAGE (Android 11+ for receiving files), CAMERA (QR Codes), NEARBY_WIFI_DEVICES (Android 13+)",
    "Other applicable policy declarations": "N/A"
}
add_table_section('5. App Content & Policy Declarations', section5)

section6 = {
    "Does the app require login? — Yes / No": "No",
    "Test username/email, if required": "N/A",
    "Test password, if required": "N/A",
    "OTP/PIN instructions, if required": "N/A",
    "Steps for the reviewer to access important features": "1. Connect two devices to the same Wi-Fi network. 2. Open ClipLAN on both. 3. Tap a discovered device on the radar to transfer files instantly.",
    "Any special instructions or setup required for review": "Must test in an environment where two devices can communicate over a local LAN/WLAN."
}
add_table_section('6. Login / Reviewer Access', section6)

doc.add_heading('Team Submission Summary', level=2)
summary_table = doc.add_table(rows=2, cols=8)
summary_table.style = 'Table Grid'
shdr = summary_table.rows[0].cells
shdr[0].text = 'Team Number'
shdr[1].text = 'Team Name'
shdr[2].text = 'Submitted By'
shdr[3].text = 'AAB'
shdr[4].text = 'Store Assets'
shdr[5].text = 'Privacy Policy'
shdr[6].text = 'Data Safety'
shdr[7].text = 'Review Access'

for cell in shdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

row1 = summary_table.rows[1].cells
row1[0].text = '-1'
row1[1].text = 'BATCH -1'
row1[2].text = 'Abdul Razeek A'
row1[3].text = 'Yes'
row1[4].text = 'Yes'
row1[5].text = 'Yes'
row1[6].text = 'Yes'
row1[7].text = 'Yes'

doc.save('ClipLAN_PlayStore_Checklist.docx')
print("ClipLAN_PlayStore_Checklist.docx generated successfully.")
