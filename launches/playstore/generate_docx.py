import os
import subprocess

try:
    import docx
except ImportError:
    subprocess.check_call(["pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Google Play Store Publishing Checklist', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Category'
hdr_cells[1].text = 'Checklist Item'
hdr_cells[2].text = 'Answer / Value'

data = [
    ('1. Basic App Details', 'Team name / team number', 'Abdul Razeek A'),
    ('', "Team members' names", 'Abdul Razeek A'),
    ('', 'App name (maximum 30 characters)', 'ClipLAN'),
    ('', 'Default app language', 'English (en-US)'),
    ('', 'App or Game', 'App'),
    ('', 'Free or Paid', 'Free'),
    ('', 'App category', 'Tools / Productivity'),
    ('', 'Tags, if applicable', 'File Sharing, Local Network, Productivity'),
    ('', 'Package name / Application ID', 'com.clipland.clipland'),
    ('', 'Contact email', 'contact@clipland.example.com'),
    
    ('2. Store Listing Details', 'Short description (maximum 80 characters)', 'Fast, serverless P2P local file sharing and shared clipboard app.'),
    ('', 'Full description (maximum 4,000 characters)', 'ClipLAN is a blazing fast, open-source, serverless, peer-to-peer Local Area Network file sharing and real-time shared clipboard utility. It requires no internet connection and transfers files securely over your local Wi-Fi router at maximum speeds.'),
    ('', 'High-quality app icon', 'Ready (launches/playstore/screenshots)'),
    ('', 'Phone screenshots', 'Ready (launches/playstore/screenshots)'),
    ('', 'Tablet screenshots, if applicable', 'N/A'),
    ('', 'Feature graphic, where required', 'Ready'),
    ('', 'Promotional video, if applicable', 'N/A'),
    ('', 'Developer/contact information', 'Abdul Razeek A'),
    
    ('3. App Build & Technical Details', 'Final Android App Bundle (.aab)', 'Ready (launches/playstore/app-release.aab)'),
    ('', 'Application/package ID matches the submitted app', 'Yes'),
    ('', 'Final production build tested', 'Yes'),
    ('', 'App launches and core features work correctly', 'Yes'),
    ('', 'No unnecessary test/demo data in the production build', 'Yes'),
    
    ('4. Privacy & Data Safety', 'Privacy Policy URL', 'https://clipland.example.com/terms.html'),
    ('', 'Does the app collect user data? — Yes / No', 'No'),
    ('', 'Does the app share user data? — Yes / No', 'No'),
    ('', 'Types of data collected', 'None'),
    ('', 'Purpose of data collection', 'N/A'),
    ('', 'Third-party SDKs/services used', 'None (Local P2P only)'),
    ('', 'Is data encrypted in transit?', 'No (Operates entirely on trusted Local LAN/Wi-Fi)'),
    ('', 'Can users request deletion of their data?', 'N/A (No data stored remotely)'),
    
    ('5. App Content & Policy Declarations', 'Contains ads? — Yes / No', 'No'),
    ('', 'Target audience / age group', 'Everyone'),
    ('', 'Content rating questionnaire information', 'Productivity / Utility'),
    ('', 'App access requirements', 'No login required'),
    ('', 'Sensitive permissions used', 'MANAGE_EXTERNAL_STORAGE (Android 11+), INTERNET (for Local Sockets)'),
    ('', 'Other applicable policy declarations', 'N/A'),
    
    ('6. Login / Reviewer Access', 'Does the app require login? — Yes / No', 'No'),
    ('', 'Test username/email, if required', 'N/A'),
    ('', 'Test password, if required', 'N/A'),
    ('', 'OTP/PIN instructions, if required', 'N/A'),
    ('', 'Steps for the reviewer to access important features', 'Install on two devices connected to the SAME Wi-Fi network. App uses UDP multicast to discover peers locally and TCP to transfer files. No internet access is required or used.'),
    ('', 'Any special instructions or setup required for review', 'Reviewers MUST use two test devices connected to the same local network to test file transfer functionality.')
]

for category, item, answer in data:
    row_cells = table.add_row().cells
    row_cells[0].text = category
    row_cells[1].text = item
    row_cells[2].text = answer

doc.add_heading('Play Store Text Assets', 0)

doc.add_heading('12. Short Description', level=1)
doc.add_paragraph('Blazing fast, serverless P2P local file sharing & real-time shared clipboard.')

doc.add_heading('13. Full Description', level=1)
doc.add_paragraph(
    "ClipLAN is a revolutionary peer-to-peer productivity tool designed to make sharing files and clipboard snippets across your devices instantaneous and effortless. Operating entirely on your Local Area Network (Wi-Fi), ClipLAN bypasses the internet entirely. This means you get maximum router speeds, zero data limits, and unparalleled privacy.\n\n"
    "Whether you are a photographer moving massive RAW files, a developer sharing APKs, or a student transferring project data, ClipLAN's optimized network engine handles it flawlessly. By utilizing UDP multicast for automatic device discovery and raw TCP sockets for data streaming, transfers begin instantly and complete in seconds.\n\n"
    "With cross-platform support and a stunning glassmorphic interface, ClipLAN is the ultimate local sharing solution."
)

doc.add_heading('14. Main Features', level=1)
features_list = [
    "🚀 Lightning-Fast Transfers: Transfer files at maximum local network speeds with zero internet bandwidth consumption.",
    "🔒 Secure & Private: Files are sent directly device-to-device. No cloud servers, no intermediate storage, and no tracking.",
    "📡 Automatic Peer Discovery: Instantly see other devices running ClipLAN on your Wi-Fi network using intelligent UDP multicast.",
    "📋 Real-Time Shared Clipboard: Copy text on your phone and instantly paste it on your computer. Your clipboards sync automatically.",
    "🔄 Cross-Platform Capability: Seamlessly share files between Android, iOS, Windows, and macOS devices.",
    "✨ Modern UI: A beautiful, responsive glassmorphic design that is intuitive and easy to use.",
    "✅ File Integrity Verification: Built-in SHA-256 hashing guarantees your files arrive exactly as they were sent without corruption.",
    "📶 No Internet Required: Works perfectly offline as long as devices are connected to the same local router or mobile hotspot."
]
for feature in features_list:
    doc.add_paragraph(feature, style='List Bullet')

doc.save('launches/playstore/PlayStore_Submission_Answers.docx')
print("Docx file generated successfully.")
